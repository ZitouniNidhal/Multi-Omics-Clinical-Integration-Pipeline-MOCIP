import re
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_bib(bib_file):
    if not os.path.exists(bib_file):
        return {}
    with open(bib_file, 'r', encoding='utf-8') as f:
        content = f.read()
    entries = {}
    for entry in re.finditer(r'@(\w+)\s*\{\s*([^,]+),', content):
        key = entry.group(2).strip()
        start = entry.end()
        end = content.find('@', start)
        if end == -1: end = len(content)
        data = content[start:end]
        
        author_match = re.search(r'author\s*=\s*[\"{](.*?)[\"}]', data, re.IGNORECASE | re.DOTALL)
        year_match = re.search(r'year\s*=\s*[\"{]?(\d+)[\"}]?', data, re.IGNORECASE)
        
        if author_match and year_match:
            authors_str = author_match.group(1).replace('{', '').replace('}', '').strip()
            names = [n.strip() for n in authors_str.split(' and ')]
            if len(names) == 1:
                last_name = names[0].split(',')[-1].strip() if ',' in names[0] else names[0].split(' ')[-1].strip()
                cite_str = f"({last_name}, {year_match.group(1)})"
            elif len(names) == 2:
                last_1 = names[0].split(',')[-1].strip() if ',' in names[0] else names[0].split(' ')[-1].strip()
                last_2 = names[1].split(',')[-1].strip() if ',' in names[1] else names[1].split(' ')[-1].strip()
                cite_str = f"({last_1} & {last_2}, {year_match.group(1)})"
            else:
                last_1 = names[0].split(',')[-1].strip() if ',' in names[0] else names[0].split(' ')[-1].strip()
                cite_str = f"({last_1} et al., {year_match.group(1)})"
            entries[key] = cite_str
    return entries

def clean_latex(text, bib_dict):
    text = re.sub(r'%.*?\n', '\n', text)
    text = re.sub(r'\\\\', '\n', text)
    
    def replace_cite(match):
        keys = match.group(1).split(',')
        cites = [bib_dict.get(k.strip(), f"({k.strip()})") for k in keys]
        return " " + "; ".join(cites).replace("((", "(").replace("))", ")")
    
    text = re.sub(r'\\cite\{(.*?)\}', replace_cite, text)
    text = re.sub(r'\\textbf\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\textit\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\texttt\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\emph\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\item', r'\n• ', text)
    
    text = text.replace('\\%', '%').replace('\\&', '&').replace('\\_', '_')
    text = text.replace('\\lbrace', '{').replace('\\rbrace', '}').replace('\\textbar', '|')
    text = text.replace('\\mid', '|').replace('\\dots', '...')
    text = text.replace('$', '')
    
    text = re.sub(r'\\begin\{.*?\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\end\{.*?\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\\([a-zA-Z]+)', '', text) 
    text = text.replace('{', '').replace('}', '')
    
    return text.strip()

def run_conversion():
    bib = parse_bib('references.bib')
    with open('multiomics_igi.tex', 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Title
    title_match = re.search(r'\\title\{(.*?)\}', content, re.DOTALL)
    if title_match:
        title_text = re.sub(r'\\([a-zA-Z]+)', '', title_match.group(1)).replace('{', '').replace('}', '').replace('\n', ' ')
        p = doc.add_heading(title_text, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Nidhal Zitouni and Mohsen Maraoui\nDepartment of Computer Science, Faculty of Sciences of Monastir, Monastir, Tunisia\nzitouninidhal@fsm.u-monastir.tn, mohsen.maraoui@fsm.rnu.tn').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', content, re.DOTALL)
    if abstract_match:
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(clean_latex(abstract_match.group(1), bib))

    # Process Sections, Subsections, and Subsubsections
    # We will use a more robust split to handle hierarchy
    parts = re.split(r'\\(section|subsection|subsubsection)\{(.*?)\}', content)
    
    # parts[0] is everything before the first section
    for i in range(1, len(parts), 3):
        level_type = parts[i]
        level_title = parts[i+1]
        level_body = parts[i+2]
        
        if 'printbibliography' in level_title or 'References' in level_title:
            break
            
        level = 1 if level_type == 'section' else (2 if level_type == 'subsection' else 3)
        doc.add_heading(level_title, level=level)
        
        # Clean and add the body of this part (until next section/subsection/subsubsection)
        # Note: level_body here contains the text until the NEXT command found by split.
        clean_text = clean_latex(level_body, bib)
        if clean_text:
            doc.add_paragraph(clean_text)

    doc.add_heading('Note on Tables and Figures', level=1)
    doc.add_paragraph('Please refer to the separate high-resolution image files and the original LaTeX source for the complex tables and diagrams. They have been omitted in this Word version as per standard IGI Global submission guidelines for typesetting.')

    doc.save('multiomics_igi.docx')
    print('Successfully created multiomics_igi.docx with full hierarchy')

if __name__ == "__main__":
    run_conversion()
